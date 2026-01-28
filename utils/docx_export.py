"""Word document export utilities for audit report assets."""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from typing import List, Tuple, Optional


# AIUC-1 brand colors
AIUC_PRIMARY = RGBColor(99, 102, 241)      # #6366f1 - Indigo
AIUC_SECONDARY = RGBColor(139, 92, 246)    # #8b5cf6 - Purple
AIUC_DARK = RGBColor(26, 26, 26)           # #1a1a1a - Surface
AIUC_BORDER = RGBColor(42, 42, 42)         # #2a2a2a - Border
AIUC_TEXT = RGBColor(255, 255, 255)        # #ffffff - White
AIUC_MUTED = RGBColor(136, 136, 136)       # #888888 - Muted

# Severity colors
SEVERITY_COLORS = {
    'PASS': RGBColor(16, 185, 129),   # #10b981 - Green
    'P4': RGBColor(251, 191, 36),     # #fbbf24 - Yellow
    'P3': RGBColor(249, 115, 22),     # #f97316 - Orange
    'P2': RGBColor(239, 68, 68),      # #ef4444 - Red
    'P1': RGBColor(220, 38, 38),      # #dc2626 - Dark Red
    'P0': RGBColor(153, 27, 27),      # #991b1b - Darkest Red
}


def rgb_to_hex(color: RGBColor) -> str:
    """Convert RGBColor to hex string (without #)."""
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'


def set_cell_background(cell, color: RGBColor):
    """Set background color of a table cell."""
    hex_color = rgb_to_hex(color)
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def style_header_row(row, bg_color: RGBColor = AIUC_PRIMARY, text_color: RGBColor = AIUC_TEXT):
    """Apply AIUC-1 styling to a header row."""
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = text_color
                run.font.size = Pt(10)
                run.font.name = 'Inter'


def style_data_row(row, row_idx: int, severity_col_idx: Optional[int] = None):
    """Apply styling to a data row with optional severity color coding."""
    # Alternate row colors for readability
    bg_color = RGBColor(250, 250, 250) if row_idx % 2 == 0 else RGBColor(255, 255, 255)

    for col_idx, cell in enumerate(row.cells):
        set_cell_background(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Inter'
                run.font.color.rgb = RGBColor(51, 51, 51)


def set_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def df_to_table(doc: Document, df: pd.DataFrame, title: str = None) -> None:
    """
    Add a DataFrame as a styled table to a Word document.

    Args:
        doc: python-docx Document object
        df: DataFrame to convert
        title: Optional section title
    """
    # Add title
    if title:
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = AIUC_PRIMARY
        heading.runs[0].font.name = 'Inter'

    # Create table
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Header row
    header_row = table.rows[0]
    for col_idx, col_name in enumerate(df.columns):
        cell = header_row.cells[col_idx]
        cell.text = str(col_name)
    style_header_row(header_row)

    # Data rows
    for row_idx, (_, row_data) in enumerate(df.iterrows()):
        table_row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = table_row.cells[col_idx]
            cell.text = str(value) if pd.notna(value) else ''
        style_data_row(table_row, row_idx)

    # Add borders
    set_table_borders(table)

    # Add spacing after table
    doc.add_paragraph()


def create_audit_report_docx(
    tables: List[Tuple[str, pd.DataFrame]],
    report_title: str = "AIUC-1 Audit Report Assets",
    subtitle: str = None
) -> BytesIO:
    """
    Create a Word document containing multiple styled tables.

    Args:
        tables: List of (title, DataFrame) tuples
        report_title: Main document title
        subtitle: Optional subtitle (e.g., date or round name)

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Document title
    title = doc.add_heading(report_title, level=0)
    title.runs[0].font.color.rgb = AIUC_PRIMARY
    title.runs[0].font.name = 'Inter'

    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.runs[0].font.color.rgb = AIUC_MUTED
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.name = 'Inter'

    doc.add_paragraph()  # Spacing

    # Add each table
    for table_title, df in tables:
        df_to_table(doc, df, title=table_title)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def create_single_table_docx(df: pd.DataFrame, title: str = None) -> BytesIO:
    """Create a Word document with a single table (for individual downloads)."""
    return create_audit_report_docx([(title, df)], report_title=title or "Table Export")
